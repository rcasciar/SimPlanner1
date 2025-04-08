"""
Modulo per l'esportazione delle programmazioni in formato Word.
"""

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import pandas as pd
from datetime import datetime, time
from models import ScheduleData
from io import BytesIO
import traceback

def create_word_document_from_schedule(schedule_data, title="Programmazione Laboratori"):
    """
    Crea un documento Word dalla programmazione dei laboratori.
    
    Args:
        schedule_data: DataFrame o ScheduleData contenente la programmazione
        title: Titolo del documento
    
    Returns:
        BytesIO contenente il documento Word
    """
    # Crea nuovo documento
    doc = docx.Document()
    
    # Configura margini delle pagine
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Aggiungi titolo
    doc.add_heading(title, level=1)
    
    # Verifica se schedule_data è un DataFrame
    if isinstance(schedule_data, pd.DataFrame):
        # Per i DataFrame, crea un documento semplificato
        days = get_days_with_data(schedule_data)
        for day_num in days:
            # Estrai la data dal formato "Giorno X"
            day_title = f"Giorno {day_num+1}"
            doc.add_heading(day_title, level=2)
            
            # Crea tabella per il giorno
            simplified_table(doc, schedule_data, day_num)
            
            # Aggiungi spazio dopo la tabella
            doc.add_paragraph("")
        
        # Restituisci il documento come BytesIO
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        return f
    
    # Aggiungi data di generazione
    current_date = datetime.now().strftime("%d/%m/%Y %H:%M")
    date_paragraph = doc.add_paragraph(f"Generato il: {current_date}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Crea tabella per ogni giorno
    days = get_days_with_data(schedule_data)
    
    for day_num in days:
        # Data di inizio (1 maggio 2025)
        start_date = datetime(2025, 5, 1)
        # Calcola la data effettiva aggiungendo il numero di giorni
        actual_date = start_date.replace(day=start_date.day + day_num)
        # Formatta la data in italiano: "GG Mese AAAA"
        day_title = actual_date.strftime("%d %B %Y").capitalize()
        doc.add_heading(day_title, level=2)
        
        # Crea tabella per il giorno
        create_day_table(doc, schedule_data, day_num)
        
        # Aggiungi spazio dopo la tabella
        doc.add_paragraph("")
    
    # Restituisci il documento come BytesIO
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    
    return f

def get_days_with_data(schedule_data):
    """
    Ottiene i giorni per cui ci sono dati di programmazione.
    
    Args:
        schedule_data: DataFrame o ScheduleData contenente la programmazione
    
    Returns:
        Lista di numeri di giorni
    """
    days = set()
    
    # Verifica se schedule_data è un DataFrame o un oggetto ScheduleData
    if isinstance(schedule_data, pd.DataFrame):
        if 'data' in schedule_data.columns:
            # Estraiamo i giorni dalle date nel DataFrame
            for data in schedule_data['data'].unique():
                if "Giorno" in data:
                    try:
                        # Estrai il numero dal formato "Giorno X"
                        day_num = int(data.split()[1]) - 1  # Converte base-1 a base-0
                        days.add(day_num)
                    except (ValueError, IndexError):
                        continue
            
        return sorted(list(days)) if days else [0]  # Default al primo giorno se non ci sono giorni identificabili
    else:
        # Se è un oggetto ScheduleData
        for lab in schedule_data.scheduled_labs:
            days.add(lab.time_slot.day)
        return sorted(list(days))

def create_day_table(doc, schedule_data, day_num):
    """
    Crea una tabella per un giorno specifico.
    
    Args:
        doc: Documento Word
        schedule_data: DataFrame o ScheduleData contenente la programmazione
        day_num: Numero del giorno (0-based)
    """
    # Verifica se schedule_data è un DataFrame o un oggetto ScheduleData
    if isinstance(schedule_data, pd.DataFrame):
        # Crea una versione semplificata per DataFrame
        simplified_table(doc, schedule_data, day_num)
        return
    
    # Se è un oggetto ScheduleData, continua con l'implementazione originale
    # Ottieni i lab schedulati per questo giorno
    day_labs = [lab for lab in schedule_data.scheduled_labs if lab.time_slot.day == day_num]
    
    if not day_labs:
        doc.add_paragraph("Nessuna programmazione per questo giorno.")
        return
    
    # Ottieni le fasce orarie
    time_slots = get_time_slots_for_day(day_labs)
    
    # Ottieni le aule
    rooms = get_rooms_for_day(day_labs)
    
    # Calcola il numero di colonne: 2 (DATA, ORARIO) + numero di aule * 2 (nome lab, studenti)
    cols = 2 + len(rooms)
    
    # Crea tabella
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    
    # Intestazioni
    header_cells = table.rows[0].cells
    header_cells[0].text = "DATA"
    header_cells[1].text = "ORARIO"
    
    for i, room in enumerate(rooms):
        header_cells[i + 2].text = room.name.upper()
    
    # Formatta header
    for cell in header_cells:
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Aggiungi righe per ogni fascia oraria
    for slot_idx, time_slot in enumerate(time_slots):
        # Prima riga: data e orario
        row = table.add_row()
        row_cells = row.cells
        
        # Data (solo sulla prima riga del giorno) in formato italiano
        if slot_idx == 0:
            # Data di inizio (1 maggio 2025)
            start_date = datetime(2025, 5, 1)
            # Calcola la data effettiva aggiungendo il numero di giorni
            actual_date = start_date.replace(day=start_date.day + day_num)
            # Formatta la data in italiano: "GG Mese AAAA"
            date_text = actual_date.strftime("%d %B %Y").capitalize()
            row_cells[0].text = date_text
        
        # Orario
        start_time = time_slot.start_time.strftime("%H.%M")
        end_time = time_slot.end_time.strftime("%H.%M")
        row_cells[1].text = f"{start_time}–{end_time}"
        
        # Per ogni aula, trova il lab programmato in questa fascia
        for room_idx, room in enumerate(rooms):
            cell = row_cells[room_idx + 2]
            scheduled_lab = find_lab_for_room_and_slot(day_labs, room.name, time_slot)
            
            if scheduled_lab:
                # Lab name
                lab_name = scheduled_lab.lab.name
                # Find student groups
                students = [f"Gruppo {student_id + 1}" for student_id in scheduled_lab.students[:8]]  # Limita a 8 per leggibilità
                if len(scheduled_lab.students) > 8:
                    students.append(f"... (+{len(scheduled_lab.students) - 8})")
                
                # Formatta il testo
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Nome del laboratorio
                lab_run = p.add_run(f"{lab_name}\n")
                lab_run.bold = True
                lab_run.font.size = Pt(9)
                
                # Durata
                duration_run = p.add_run(f"({scheduled_lab.lab.duration_minutes/60:.1f} ore)\n")
                duration_run.font.size = Pt(8)
                duration_run.italic = True
                
                # Gruppi di studenti
                for student in students:
                    student_run = p.add_run(f"{student}\n")
                    student_run.font.size = Pt(9)
        
        # Seconda riga: tutor
        tutor_row = table.add_row()
        tutor_cells = tutor_row.cells
        
        # Celle vuote per data e orario
        tutor_cells[0].text = ""
        tutor_cells[1].text = ""
        
        # Per ogni aula, aggiungi il tutor (simulato)
        for room_idx, room in enumerate(rooms):
            tutor_cells[room_idx + 2].text = "TUTOR"
            p = tutor_cells[room_idx + 2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.runs[0]
            run.font.size = Pt(8)
            run.italic = True
    
    # Aggiungi riga per la pausa pranzo se necessario
    if any(ts.start_time.hour >= 13 for ts in time_slots):
        lunch_idx = next((i for i, ts in enumerate(time_slots) if ts.start_time.hour >= 13), None)
        
        if lunch_idx is not None:
            # Inserisci la riga della pausa pranzo prima della fascia oraria post-pranzo
            lunch_row = table.add_row()
            lunch_cells = lunch_row.cells
            
            # Formatta la riga della pausa pranzo
            lunch_cells[0].text = ""
            lunch_cells[1].text = "13.10–14.10"
            
            # Imposta "PAUSA PRANZO" su tutte le altre celle
            for i in range(2, cols):
                lunch_cells[i].text = "PAUSA PRANZO"
                lunch_paragraph = lunch_cells[i].paragraphs[0]
                lunch_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lunch_run = lunch_paragraph.runs[0]
                lunch_run.bold = True
                lunch_run.font.size = Pt(9)
    
    # Aggusta larghezza colonne
    table.autofit = False
    table.columns[0].width = Cm(2.0)  # DATA
    table.columns[1].width = Cm(2.5)  # ORARIO
    
    # Distribuisci lo spazio rimanente tra le aule
    remaining_width = Cm(16)  # Larghezza totale stimata della pagina meno le prime due colonne
    room_width = remaining_width / len(rooms)
    for i in range(2, cols):
        table.columns[i].width = room_width

def get_time_slots_for_day(day_labs):
    """
    Ottiene le fasce orarie uniche per un giorno specifico.
    
    Args:
        day_labs: Lista di ScheduledLab per il giorno
    
    Returns:
        Lista di TimeSlot ordinate per orario
    """
    # Estrai le fasce orarie uniche
    time_slots = set()
    for lab in day_labs:
        time_slots.add(lab.time_slot)
    
    # Converti in lista e ordina per orario
    return sorted(list(time_slots), key=lambda x: x.start_time)

def get_rooms_for_day(day_labs):
    """
    Ottiene le aule uniche per un giorno specifico.
    
    Args:
        day_labs: Lista di ScheduledLab per il giorno
    
    Returns:
        Lista di Room
    """
    # Estrai le aule uniche
    rooms = set()
    for lab in day_labs:
        rooms.add(lab.room)
    
    # Converti in lista
    return list(rooms)

def find_lab_for_room_and_slot(day_labs, room_name, time_slot):
    """
    Trova il laboratorio programmato per una specifica aula e fascia oraria.
    
    Args:
        day_labs: Lista di ScheduledLab per il giorno
        room_name: Nome dell'aula
        time_slot: Fascia oraria
    
    Returns:
        ScheduledLab o None se non trovato
    """
    for lab in day_labs:
        if lab.room.name == room_name and lab.time_slot == time_slot:
            return lab
    
    return None


def simplified_table(doc, df, day_num):
    """
    Crea una tabella semplificata per DataFrame.
    
    Args:
        doc: Documento Word
        df: DataFrame con la programmazione
        day_num: Numero del giorno (0-based)
    """
    # Filtra il DataFrame per il giorno specifico
    day_string = f"Giorno {day_num+1}"
    day_df = df[df['data'] == day_string] if 'data' in df.columns else df
    
    if day_df.empty:
        doc.add_paragraph("Nessuna programmazione per questo giorno.")
        return
    
    # Crea tabella con tutte le colonne del DataFrame
    num_cols = len(day_df.columns)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Intestazioni
    for i, col_name in enumerate(day_df.columns):
        table.cell(0, i).text = col_name.upper()
        cell_paragraph = table.cell(0, i).paragraphs[0]
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(10)
    
    # Dati
    for i, row in day_df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(day_df.columns):
            row_cells[j].text = str(row[col])
    
    # Impostazioni di larghezza colonne
    table.autofit = True


def export_schedule_word(schedule_data, filename="programmazione_laboratori.docx", sede_cdl=None, anno_corso=None, num_macrogruppi=None):
    """
    Esporta la programmazione dei laboratori in un file Word.
    
    Args:
        schedule_data: DataFrame o ScheduleData contenente la programmazione
        filename: Nome del file Word
        sede_cdl: Sede del Corso di Laurea
        anno_corso: Anno di corso
        num_macrogruppi: Numero di macrogruppi (canali)
    
    Returns:
        BytesIO contenente il file Word
    """
    try:
        # Log dettagliato per debug
        print(f"Tipo di schedule_data: {type(schedule_data)}")
        
        if isinstance(schedule_data, pd.DataFrame):
            print(f"Colonne nel DataFrame: {schedule_data.columns.tolist()}")
            print(f"Numero di righe: {len(schedule_data)}")
        
        # Crea documento Word
        doc_or_buffer = create_word_document_from_schedule(schedule_data, title="Programmazione Laboratori")
        
        # Gestiamo il caso in cui doc sia già un BytesIO
        if isinstance(doc_or_buffer, BytesIO):
            # Se è già un BytesIO, probabilmente è già stato salvato
            print("Documento Word già generato come BytesIO")
            doc_or_buffer.seek(0)
            return doc_or_buffer
        
        doc = doc_or_buffer
        
        # Aggiungi informazioni sulla sede e anno di corso, se presenti
        if sede_cdl or anno_corso or num_macrogruppi:
            # Trova il primo paragrafo (titolo)
            first_paragraph = None
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    first_paragraph = paragraph
                    break
            
            if first_paragraph:
                # Aggiungi un paragrafo dopo il titolo con sede e anno
                header_info = []
                if sede_cdl:
                    header_info.append(f"Sede: {sede_cdl}")
                if anno_corso:
                    header_info.append(f"Anno di corso: {anno_corso}")
                if num_macrogruppi and num_macrogruppi > 1:
                    canali = ", ".join([f"Canale {chr(65+i)}" for i in range(num_macrogruppi)])
                    header_info.append(f"Canali: {canali}")
                
                # Inserisci paragrafo dopo il titolo
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(" | ".join(header_info)).font.size = Pt(10)
                
                # Sposta il paragrafo appena dopo il titolo
                try:
                    for i in range(len(doc.paragraphs) - 1, 0, -1):
                        if doc.paragraphs[i] == first_paragraph:
                            # Sposta il paragrafo dopo il titolo
                            doc._element.body.insert(i + 1, p._element)
                            # Rimuovi il paragrafo aggiunto alla fine
                            doc._element.body.remove(doc.paragraphs[-1]._element)
                            break
                except Exception as e:
                    print(f"Errore nella manipolazione dei paragrafi: {str(e)}")
                    # Non interrompere l'esportazione per questo errore
        
        # Salva in BytesIO
        buffer = BytesIO()
        try:
            doc.save(buffer)
            buffer.seek(0)
            print("Documento Word salvato con successo")
            return buffer
        except Exception as e:
            print(f"Errore nel salvataggio del documento Word: {str(e)}")
            # Fallback: prova a creare un nuovo documento minimo
            try:
                fallback_doc = docx.Document()
                fallback_doc.add_heading("Programmazione Laboratori", 0)
                fallback_doc.add_paragraph("Errore nella generazione del documento completo. Si consiglia di utilizzare l'export Excel.")
                fallback_buffer = BytesIO()
                fallback_doc.save(fallback_buffer)
                fallback_buffer.seek(0)
                print("Generato documento fallback dopo errore")
                return fallback_buffer
            except Exception as e2:
                print(f"Anche il fallback è fallito: {str(e2)}")
                return None
    
    except Exception as outer_e:
        # Cattura qualsiasi errore nella funzione principale
        error_trace = traceback.format_exc()
        print(f"Errore in export_schedule_word: {str(outer_e)}")
        print(f"Traceback: {error_trace}")
        
        # Prova un ultimo fallback
        try:
            minimal_doc = docx.Document()
            minimal_doc.add_heading("Errore", 0)
            minimal_doc.add_paragraph(f"Si è verificato un errore: {str(outer_e)}")
            minimal_buffer = BytesIO()
            minimal_doc.save(minimal_buffer)
            minimal_buffer.seek(0)
            return minimal_buffer
        except:
            return None