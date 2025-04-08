import docx
from pprint import pprint

# Carica il documento docx
doc = docx.Document("attached_assets/PROGRAMMAZIONE_TIRLAB_ASL_CDT_maggio2025_rev20250312_per tutor.docx")

print(f"Il documento contiene {len(doc.tables)} tabelle.")

# Analizza la prima tabella (presumibilmente la tabella di programmazione)
print("\nAnalizziamo la prima tabella:")
table = doc.tables[0]
num_rows = len(table.rows)
num_cols = len(table.rows[0].cells)
print(f"Righe: {num_rows}, Colonne: {num_cols}")

# Stampa le intestazioni
print("\nIntestazioni:")
headers = [cell.text for cell in table.rows[0].cells]
print(headers)

# Stampa le prime 3 righe per vedere la struttura
print("\nPrime 3 righe di contenuto:")
for i in range(1, min(4, num_rows)):
    row_data = [cell.text for cell in table.rows[i].cells]
    print(f"Riga {i}: {row_data}")

# Analizza anche la seconda tabella se presente
print("\nAnalisi di tutte le tabelle:")

# Analizza tutte le tabelle per trovare la pausa pranzo
for t_idx, table in enumerate(doc.tables):
    if t_idx > 2:  # Limitiamo a 3 tabelle
        break
        
    print(f"\nTabella {t_idx}:")
    num_rows = len(table.rows)
    if num_rows == 0:
        print("  Tabella vuota")
        continue
        
    num_cols = len(table.rows[0].cells)
    print(f"  Righe: {num_rows}, Colonne: {num_cols}")
    
    # Cerca la pausa pranzo nella colonna dell'orario
    orario_col = 1  # La colonna ORARIO sembra essere la seconda (indice 1)
    pause_rows = []
    
    for i in range(num_rows):
        row_data = [cell.text for cell in table.rows[i].cells]
        orario = row_data[orario_col] if len(row_data) > orario_col else ""
        # Cerca testo che contiene "pranzo" o "PAUSA"
        if "pranzo" in orario.lower() or "pausa" in orario.upper():
            pause_rows.append(i)
            print(f"  Trovata pausa pranzo in riga {i}: {orario}")
        
        # Stampa comunque tutte le righe per vedere la struttura
        print(f"  Riga {i}: {row_data}")